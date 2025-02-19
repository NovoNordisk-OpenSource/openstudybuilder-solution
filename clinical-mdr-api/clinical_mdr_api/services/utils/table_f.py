import os
from collections import defaultdict
from typing import Annotated, Any, Mapping

import yattag
from docx.shared import Inches
from openpyxl import Workbook, load_workbook
from openpyxl.styles import NamedStyle
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo
from openpyxl.worksheet.worksheet import Worksheet
from pydantic import BaseModel, Field

from clinical_mdr_api.services.utils.docx_builder import DocxBuilder
from common.telemetry import trace_calls

CHAR_WIDTHS = {
    "i": 0.5,
    "l": 0.5,
    "t": 0.6,
    "r": 0.6,
    " ": 0.5,
    "w": 1.5,
    "m": 1.5,
    "f": 0.7,
    "j": 0.7,
    "k": 0.7,
    "a": 1.0,
    "b": 1.0,
    "c": 1.0,
    "d": 1.0,
    "e": 1.0,
    "g": 1.0,
    "h": 1.0,
    "n": 1.0,
    "o": 1.0,
    "p": 1.0,
    "q": 1.0,
    "s": 1.0,
    "u": 1.0,
    "v": 1.0,
    "x": 1.0,
    "y": 1.0,
    "z": 1.0,
    "A": 1.2,
    "B": 1.2,
    "C": 1.2,
    "D": 1.2,
    "E": 1.2,
    "F": 1.2,
    "G": 1.2,
    "H": 1.2,
    "I": 0.6,
    "J": 1.2,
    "K": 1.2,
    "L": 1.2,
    "M": 1.5,
    "N": 1.2,
    "O": 1.2,
    "P": 1.2,
    "Q": 1.2,
    "R": 1.2,
    "S": 1.2,
    "T": 1.2,
    "U": 1.2,
    "V": 1.2,
    "W": 1.5,
    "X": 1.2,
    "Y": 1.2,
    "Z": 1.2,
    "0": 1.0,
    "1": 0.8,
    "2": 1.0,
    "3": 1.0,
    "4": 1.0,
    "5": 1.0,
    "6": 1.0,
    "7": 1.0,
    "8": 1.0,
    "9": 1.0,
    ".": 0.5,
    ",": 0.5,
    ";": 0.5,
    ":": 0.5,
    "!": 0.5,
    "?": 0.5,
    "-": 0.5,
    "_": 0.5,
    "(": 0.6,
    ")": 0.6,
    "[": 0.6,
    "]": 0.6,
    "{": 0.6,
    "}": 0.6,
    "/": 0.6,
    "\\": 0.6,
    "|": 0.4,
    "'": 0.3,
    '"': 0.4,
}


class Ref(BaseModel):
    type: Annotated[str | None, Field(title="Referenced item type")]
    uid: Annotated[str, Field(title="Referenced item uid")]

    class Config:
        # make the model immutable & hashable
        frozen = True

    def __init__(self, type_=None, uid=None, **kwargs):
        if type_ is not None:
            kwargs["type"] = type_
        if uid is not None:
            kwargs["uid"] = uid
        super().__init__(**kwargs)


class TableCell(BaseModel):
    text: Annotated[str, Field(title="Text contents of cell")] = ""
    span: Annotated[int, Field(title="Horizontal spanning of cell, 1 by default")] = 1
    style: Annotated[
        str | None, Field(title="Associated style to cell", nullable=True)
    ] = None
    refs: Annotated[
        list[Ref] | None, Field(title="Reference to item", nullable=True)
    ] = None
    footnotes: Annotated[
        list[str] | None, Field(title="Referenced footnotes", nullable=True)
    ] = None
    vertical: Annotated[
        bool | None, Field(title="Text text direction", nullable=True)
    ] = None

    def __init__(self, text=None, **kwargs):
        if text is not None:
            kwargs["text"] = text
        super().__init__(**kwargs)


class TableRow(BaseModel):
    cells: Annotated[
        list[TableCell], Field(default_factory=list, title="Table cells in the row")
    ]
    hide: Annotated[bool, Field(title="Hide row from display")] = False

    def __init__(self, cells=None, **kwargs):
        if cells is not None:
            kwargs["cells"] = cells
        super().__init__(**kwargs)


class SimpleFootnote(BaseModel):
    uid: Annotated[str, Field(title="StudySoAFootnote.uid")]
    text_html: Annotated[str, Field(title="HTML text of footnote")]
    text_plain: Annotated[str, Field(title="Plain text of footnote")]


class TableWithFootnotes(BaseModel):
    rows: Annotated[
        list[TableRow], Field(default_factory=list, title="List of table rows")
    ]
    footnotes: Annotated[
        dict[str, SimpleFootnote] | None,
        Field(title="Mapping of symbols and table footnotes", nullable=True),
    ] = None
    num_header_rows: Annotated[int, Field(title="Number of header rows")] = 0
    num_header_cols: Annotated[int, Field(title="Number of header columns")] = 0
    title: Annotated[
        str | None, Field(title="Table title (when rendered to HTML)", nullable=True)
    ] = None
    id: Annotated[
        str | None, Field(title="Table id (when rendered to HTML)", nullable=True)
    ] = None


@trace_calls()
def table_to_docx(
    table: TableWithFootnotes,
    styles: Mapping[str, tuple[str, Any]] | None = None,
    template: str | None = None,
) -> DocxBuilder:
    # assume horizontal table dimension from number of cells in first row
    num_cols = sum((c.span for c in table.rows[0].cells))

    # parses an empty template DOCX file into a helper class
    docx = DocxBuilder(
        styles=styles, landscape=True, margins=[0.5, 0.5, 0.5, 0.5], template=template
    )

    # adds a table to the document
    x_table = docx.create_table(
        num_rows=sum(1 for row in table.rows if not row.hide),
        num_columns=num_cols,
    )

    # set width of first column
    x_table.columns[0].width = Inches(4)

    # cache porperty for performance issues, see github.com/python-openxml/python-docx/issues/174
    x_cells = x_table._cells
    x_rows = x_table.rows

    for r, t_row in enumerate((row for row in table.rows if not row.hide)):
        num_merge, merge_to = 0, None

        # set header row to repeat on each page
        if r < table.num_header_rows:
            docx.repeat_table_header(x_rows[r])

        for c, t_cell in enumerate(t_row.cells):
            x_cell = x_cells[r * num_cols + c]

            # merge with previous spanning cell
            if num_merge:
                merge_to.merge(x_cell)
                num_merge -= 1
                continue

            # skip invisible cells (should not get here if spans are coherent)
            if t_cell.span < 1:
                continue

            # when cell span > 1 merge the following N cells into this one
            num_merge = t_cell.span - 1
            merge_to = x_cell

            # all docx cells host one or more paragraph for contents
            x_para = x_cell.paragraphs[0]

            # set cell text in the paragraph
            if t_cell.text:
                x_para.text = t_cell.text

            # resolve style name and apply to paragraph
            style_name = styles.get(t_cell.style, [None])[0] if styles else None
            if style_name:
                x_para.style = style_name

            # set vertical text direction
            if t_cell.vertical:
                docx.set_vertical_cell_direction(x_cell, "btLr")

            # add a new run (like <span>) within the paragraph for each footnote symbol, with spacing in the run
            for symbol in t_cell.footnotes or []:
                run = x_para.add_run(f" {symbol}")
                run.font.superscript = True

    # add footnotes
    if table.footnotes:
        style_name = styles.get("footnote", [None])[0] if styles else None

        for symbol, footnote in table.footnotes.items():
            # each footnote is a new paragraph at the end of the document
            x_para = docx.document.add_paragraph(style=style_name)

            # footnote symbols into a run (like <span>) with superscript
            run = x_para.add_run(symbol)
            run.font.superscript = True

            # footnote text with glue and spacing into a distinct run
            x_para.add_run(f": {footnote.text_plain}")

    return docx


@trace_calls
def table_to_html(table: TableWithFootnotes, css_style: str | None = None) -> str:
    """Renders TableWithFootnotes into an HTML document

    Renders TableWithFootnotes into an HTML document with a TABLE and footnotes into a DL (if they exist).
    Optional CSS text can be provided in `css_style` added as <style> tag.

    :param table: The table data to be rendered into HTML, including rows, cells, headers, and footnotes.
    :type table: TableWithFootnotes
    :param css_style: CSS text to be added as <style type="text/css"> tag in the HTML head.
    :type css_style: str
    :return: The rendered HTML document as a string.
    :rtype: str
    """

    doc, tag, text, line = yattag.Doc().ttl()
    doc.asis("<!DOCTYPE html>")

    with tag("html", lang="en"):
        with tag("head"):
            if table.title:
                line("title", table.title)
            if css_style:
                with tag("style", type="text/css"):
                    text(css_style)

        with tag("body"):
            attrs = {"id": table.id} if table.id else {}
            with tag("table", **attrs):
                with tag("thead"):
                    for row in table.rows[: table.num_header_rows]:
                        if row.hide:
                            continue

                        with tag("tr"):
                            for cell in row.cells:
                                if cell.span == 0:
                                    continue

                                with tag("th", **_cell_to_attrs(cell)):
                                    text(cell.text)
                                    for symbol in cell.footnotes or []:
                                        doc.asis("&nbsp;")
                                        line("sup", symbol)

                with tag("tbody"):
                    for row in table.rows[table.num_header_rows :]:
                        if row.hide:
                            continue

                        with tag("tr"):
                            for i, cell in enumerate(row.cells):
                                if cell.span == 0:
                                    continue

                                with tag(
                                    ("th" if i < table.num_header_cols else "td"),
                                    **_cell_to_attrs(cell),
                                ):
                                    text(cell.text)
                                    for symbol in cell.footnotes or []:
                                        doc.asis("&nbsp;")
                                        line("sup", symbol)

            if table.footnotes:
                with tag("dl", klass="footnotes"):
                    for symbol, footnote in table.footnotes.items():
                        line("dt", symbol)
                        with tag("dd"):
                            text(footnote.text_plain)

    return yattag.indent(doc.getvalue())


def _cell_to_attrs(cell):
    attrs = {}

    if cell.style:
        attrs["klass"] = cell.style

    if cell.span > 1:
        attrs["colspan"] = cell.span

    return attrs


@trace_calls
def table_to_xlsx(
    table: TableWithFootnotes,
    styles: Mapping[str, str] | None = None,
    template: str | None = None,
) -> Workbook:
    if template:
        template = os.path.join(os.path.dirname(__file__), template)
        workbook = load_workbook(template)

    else:
        workbook = Workbook()

    worksheet: Worksheet = workbook.active

    if table.title:
        worksheet.title = table.title

    column_widths = defaultdict(list)
    for r, row in enumerate(table.rows, start=1):
        worksheet.append([cell.text for cell in row.cells])

        for c, cell in enumerate(row.cells, start=1):
            if cell.span > 1:
                worksheet.merge_cells(
                    start_row=r,
                    start_column=c,
                    end_row=r,
                    end_column=(c + cell.span - 1),
                )

            if cell.span == 1:
                # take into account for column width calculations
                column_widths[c].append(estimate_string_length(cell.text))

    # calculate and set column widths
    for c, widths in column_widths.items():
        worksheet.column_dimensions[get_column_letter(c)].width = max(
            2, int(round(max(widths) * 1.05 + 1))
        )

    # apply named styles on cells
    if styles:
        for style_name in styles.values():
            if style_name not in workbook.named_styles:
                workbook.add_named_style(NamedStyle(style_name))

        for r, row in enumerate(worksheet.iter_rows()):
            for c, cell in enumerate(row):
                if (
                    c < len(table.rows[r].cells)
                    and table.rows[r].cells[c].style in styles
                ):
                    cell.style = styles[table.rows[r].cells[c].style]

    # define table
    tab = Table(
        displayName="Table1",
        ref=f"A1:{get_column_letter(len(table.rows[-1].cells))}{len(table.rows)}",
    )
    tab.tableStyleInfo = TableStyleInfo(name="TableStyleMedium2")

    # freeze header rows and columns
    worksheet.freeze_panes = (
        f"{get_column_letter(table.num_header_cols+1)}{table.num_header_rows+1}"
    )

    return workbook


def estimate_string_length(string: str) -> float:
    return sum((CHAR_WIDTHS.get(c, 1) for c in string))
