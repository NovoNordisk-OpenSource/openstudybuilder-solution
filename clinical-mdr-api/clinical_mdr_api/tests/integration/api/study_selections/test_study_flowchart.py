# pylint: disable=redefined-outer-name,unused-argument
import csv
from copy import deepcopy
from io import BytesIO

import docx
import docx.enum.section
import docx.enum.text
import docx.section
import docx.table
import docx.text.paragraph
import lxml.etree
import openpyxl
import pytest
from bs4 import BeautifulSoup
from starlette.testclient import TestClient

from clinical_mdr_api.domain_repositories.study_selections.study_soa_repository import (
    SoALayout,
)
from clinical_mdr_api.models.study_selections.study import (
    StatusChangeDescription,
    StudyDescriptionJsonModel,
    StudyMetadataJsonModel,
    StudyPatchRequestJsonModel,
)
from clinical_mdr_api.services.studies.study import StudyService
from clinical_mdr_api.services.studies.study_flowchart import (
    DOCX_STYLES,
    OPERATIONAL_DOCX_STYLES,
    StudyFlowchartService,
)
from clinical_mdr_api.services.utils.table_f import TableWithFootnotes
from clinical_mdr_api.tests.fixtures.database import TempDatabasePopulated

# pylint: disable=unused-import
from clinical_mdr_api.tests.integration.services.test_study_flowchart import (
    detailed_soa_table__days,
    detailed_soa_table__weeks,
    operational_soa_table__days,
    operational_soa_table__weeks,
    protocol_soa_table__days,
    protocol_soa_table__weeks,
)
from clinical_mdr_api.tests.integration.utils.factory_soa import SoATestData
from clinical_mdr_api.tests.integration.utils.utils import TestUtils
from clinical_mdr_api.tests.unit.models.test_table_f import (
    compare_docx_footnotes,
    compare_docx_table,
    compare_html_footnotes,
    compare_html_table,
)
from clinical_mdr_api.tests.utils.checks import (
    DOCX_CONTENT_TYPE,
    assert_json_response,
    assert_response_content_type,
    assert_response_status_code,
)

PROTOCOL_SOA_EXPORT_COLUMN_HEADERS = [
    "study_number",
    "study_version",
    "soa_group",
    "activity_group",
    "activity_subgroup",
    "visit",
    "activity",
]
PROTOCOL_SOA_JSON_EXPORT_COLUMN_HEADERS = PROTOCOL_SOA_EXPORT_COLUMN_HEADERS + [
    "study_activity_schedule_uid"
]
DETAILED_SOA_EXPORT_COLUMN_HEADERS = PROTOCOL_SOA_EXPORT_COLUMN_HEADERS + [
    "is_data_collected"
]
DETAILED_SOA_JSON_EXPORT_COLUMN_HEADERS = DETAILED_SOA_EXPORT_COLUMN_HEADERS + [
    "study_activity_schedule_uid"
]
OPERATIONAL_SOA_EXPORT_COLUMN_HEADERS = [
    "study_number",
    "study_version",
    "soa_group",
    "activity_group",
    "activity_subgroup",
    "epoch",
    "visit",
    "activity",
    "activity_instance",
    "topic_code",
    "param_code",
]
OPERATIONAL_SOA_JSON_EXPORT_COLUMN_HEADERS = OPERATIONAL_SOA_EXPORT_COLUMN_HEADERS + [
    "study_activity_schedule_uid"
]
OPERATIONAL_SOA_EXPORT_COLUMN_HEADERS_XLSX = [
    "Study number",
    "Study version",
    "SoA group",
    "Activity group",
    "Activity subgroup",
    "Epoch",
    "Visit",
    "Activity",
    "Activity instance",
    "Topic code",
    "Param code",
]


@pytest.fixture(scope="module")
def soa_test_data(temp_database_populated: TempDatabasePopulated) -> SoATestData:
    return SoATestData(project=temp_database_populated.project)


@pytest.mark.parametrize(
    "time_unit, layout",
    [
        (None, SoALayout.DETAILED),
        (None, SoALayout.DETAILED),
        (None, SoALayout.OPERATIONAL),
        ("day", SoALayout.DETAILED),
        ("day", SoALayout.DETAILED),
        ("day", SoALayout.OPERATIONAL),
        ("week", SoALayout.DETAILED),
        ("week", SoALayout.DETAILED),
        ("week", SoALayout.OPERATIONAL),
        (None, SoALayout.PROTOCOL),
    ],
)
def test_flowchart(
    soa_test_data: SoATestData,
    api_client: TestClient,
    time_unit: str,
    layout: SoALayout,
):
    """Test /studies/{study_uid}/flowchart returns a valid TableWithFootnotes object as JSON"""

    params = {"layout": layout.value}
    if time_unit is not None:
        params["time_unit"] = time_unit

    response = api_client.get(
        f"/studies/{soa_test_data.study.uid}/flowchart", params=params
    )
    assert_response_status_code(response, 200)
    assert_json_response(response)
    data = response.json()
    assert isinstance(data, dict)
    assert TableWithFootnotes(**data)


def test_flowchart_study_versioning(soa_test_data: SoATestData, api_client):
    locked_study_version = str(
        StudyService()
        .lock(uid=soa_test_data.study.uid, change_description="v1")
        .current_metadata.version_metadata.version_number
    )
    StudyService().unlock(uid=soa_test_data.study.uid)

    soa_test_data.create_study_visits(
        {
            "V79": {
                "epoch": "Follow-Up",
                "visit_contact_mode": "Virtual Visit",
                "day": 79,
                "min_window": -3,
                "max_window": 3,
            }
        }
    )

    response = api_client.get(
        f"/studies/{soa_test_data.study.uid}/flowchart",
        params={
            "study_value_version": locked_study_version,
            "layout": SoALayout.OPERATIONAL.value,
        },
    )
    assert_response_status_code(response, 200)
    assert_json_response(response)
    locked_data = response.json()

    response = api_client.get(
        f"/studies/{soa_test_data.study.uid}/flowchart",
        params={
            "layout": SoALayout.OPERATIONAL.value,
        },
    )
    assert_response_status_code(response, 200)
    assert_json_response(response)
    recent_data = response.json()
    assert recent_data != locked_data
    assert len(recent_data["rows"][1]["cells"]) > len(locked_data["rows"][1]["cells"])


def test_flowchart_with_non_latest_activities(soa_test_data: SoATestData, api_client):
    for activity in soa_test_data.activities:
        text_value_2_name = soa_test_data.activities[activity].name + " new version"
        # change activity name and approve the version
        response = api_client.post(
            f"/concepts/activities/activities/{soa_test_data.activities[activity].uid}/versions",
        )
        assert_response_status_code(response, 201)
        response = api_client.patch(
            f"/concepts/activities/activities/{soa_test_data.activities[activity].uid}",
            json={
                "change_description": "Change to have a new version not updated in library",
                "name": text_value_2_name,
                "name_sentence_case": text_value_2_name,
                "guidance_text": "don't know",
            },
        )
        assert_response_status_code(response, 200)
        response = api_client.post(
            f"/concepts/activities/activities/{soa_test_data.activities[activity].uid}/approvals"
        )
        assert_response_status_code(response, 201)

    response = api_client.get(
        f"/studies/{soa_test_data.study.uid}/flowchart",
    )
    assert_response_status_code(response, 200)
    assert_json_response(response)


def test_flowchart_coordinates(soa_test_data: SoATestData, api_client):
    """Tests /studies/{study_uid}/flowchart/coordinates to return a dict of coordinates (array of 2)"""

    response = api_client.get(
        f"/studies/{soa_test_data.study.uid}/flowchart/coordinates"
    )
    assert_response_status_code(response, 200)
    assert_json_response(response)
    results: dict[str, list[int, int]] = response.json()
    assert isinstance(results, dict)
    assert len(results), "no coordinates returned"


@pytest.mark.parametrize(
    "soa_table, layout, time_unit",
    [
        ("protocol_soa_table__days", SoALayout.PROTOCOL, "day"),
        ("protocol_soa_table__weeks", SoALayout.PROTOCOL, "week"),
        ("protocol_soa_table__days", SoALayout.PROTOCOL, "day"),
        ("protocol_soa_table__weeks", SoALayout.PROTOCOL, "week"),
        ("protocol_soa_table__weeks", SoALayout.PROTOCOL, "week"),
        ("detailed_soa_table__days", SoALayout.DETAILED, "day"),
        ("detailed_soa_table__weeks", SoALayout.DETAILED, "week"),
        ("operational_soa_table__days", SoALayout.OPERATIONAL, "day"),
        ("operational_soa_table__weeks", SoALayout.OPERATIONAL, "week"),
        ("operational_soa_table__days", SoALayout.OPERATIONAL, "day"),
    ],
)
def test_flowchart_html(
    request,
    soa_test_data: SoATestData,
    api_client: TestClient,
    soa_table: TableWithFootnotes,
    layout: SoALayout,
    time_unit: str,
):
    """
    Tests HTML SoA by comparing it with JSON SoA

    Go fix test_flowchart() first if both tests are failing.
    """

    soa_table: TableWithFootnotes = deepcopy(request.getfixturevalue(soa_table))

    # Layout alterations of get_study_flowchart_docx()
    if layout != SoALayout.PROTOCOL:
        StudyFlowchartService.show_hidden_rows(soa_table.rows)

    # Query parameters
    params = {"layout": layout.value}
    if time_unit is not None:
        params["time_unit"] = time_unit

    # Fetch HTML SoA
    response = api_client.get(
        f"/studies/{soa_test_data.study.uid}/flowchart.html", params=params
    )
    assert_response_status_code(response, 200)
    # THEN returns HTML content-type HTTP header
    assert_response_content_type(response, "text/html")

    # THEN body is a parsable HTML document
    # html.parser doesn't supplement the document with additional elements for the sake of well-formedness
    doc = BeautifulSoup(response.text, features="html.parser")
    # THEN document has HTML tag
    assert doc.find("html"), "<HTML> tag not found"

    # THEN the document has HTML > HEAD > TITLE tags
    assert (title := doc.select_one("html > head > title")), "<TITLE> tag not found"
    # THEN document title matches
    assert title.text == soa_table.title, "title mismatch"

    # THEN the document has TABLE
    assert (table := doc.select_one("html > body > table")), "<TABLE> tag not found"

    # Although table_f.table_to_html() has it's unit tests, we also run them on this SoA table
    # to increase the number of cases and to test on real-world scenarios.

    if layout != SoALayout.PROTOCOL:
        # Detailed and Operation SoA show all rows
        StudyFlowchartService.show_hidden_rows(soa_table.rows)

    # Compares table rows and cell contents and formatting
    compare_html_table(table, soa_table)

    if layout != SoALayout.OPERATIONAL:
        # Compares the footnote listing
        compare_html_footnotes(doc, soa_table)


@pytest.mark.parametrize(
    "soa_table, layout, time_unit",
    [
        ("protocol_soa_table__days", SoALayout.PROTOCOL, "day"),
        ("protocol_soa_table__weeks", SoALayout.PROTOCOL, "week"),
        ("protocol_soa_table__days", SoALayout.PROTOCOL, "day"),
        ("protocol_soa_table__weeks", SoALayout.PROTOCOL, "week"),
        ("protocol_soa_table__weeks", SoALayout.PROTOCOL, "week"),
        ("detailed_soa_table__days", SoALayout.DETAILED, "day"),
        ("detailed_soa_table__weeks", SoALayout.DETAILED, "week"),
        ("operational_soa_table__days", SoALayout.OPERATIONAL, "day"),
        ("operational_soa_table__weeks", SoALayout.OPERATIONAL, "week"),
        ("operational_soa_table__days", SoALayout.OPERATIONAL, "day"),
    ],
)
def test_flowchart_docx(
    request,
    soa_test_data: SoATestData,
    api_client: TestClient,
    soa_table: TableWithFootnotes,
    layout: SoALayout,
    time_unit: str,
):
    """
    Tests DOCX SoA by comparing it with JSON SoA

    Go fix test_flowchart() first if both tests are failing.
    """
    soa_table: TableWithFootnotes = deepcopy(request.getfixturevalue(soa_table))

    # Layout alterations of get_study_flowchart_docx()
    if layout != SoALayout.PROTOCOL:
        StudyFlowchartService.show_hidden_rows(soa_table.rows)

    if layout == SoALayout.PROTOCOL:
        StudyFlowchartService.add_protocol_section_column(soa_table)

    # Query parameters
    params = {"layout": layout.value}
    if time_unit is not None:
        params["time_unit"] = time_unit

    # Fetch DOCX SoA
    response = api_client.get(
        f"/studies/{soa_test_data.study.uid}/flowchart.docx", params=params
    )
    assert_response_status_code(response, 200)
    # THEN returns DOCX content type HTTP header
    assert_response_content_type(response, DOCX_CONTENT_TYPE)

    # THEN body is a parsable DOCX document
    docx_doc = docx.Document(BytesIO(response.content))

    # THEN the document contains exactly one table
    assert len(docx_doc.tables) == 1, "expected exactly 1 table in DOCX SoA"

    # Although table_f.table_to_docx() has it's unit test, we also run that on this SoA table
    # to increase the number of cases and to test on real-world scenarios.

    # Compare table rows and column contents and properties
    compare_docx_table(
        docx_doc.tables[0],
        soa_table,
        OPERATIONAL_DOCX_STYLES if layout == SoALayout.OPERATIONAL else DOCX_STYLES,
    )

    if layout != SoALayout.OPERATIONAL:
        # Compares footnote listing
        compare_docx_footnotes(
            docx_doc,
            soa_table.footnotes,
            DOCX_STYLES,
        )


@pytest.mark.parametrize(
    "time_unit, layout, debug_uids, debug_coordinates, debug_propagation",
    [
        ("day", SoALayout.DETAILED, True, None, None),
        ("day", SoALayout.DETAILED, None, True, None),
        ("day", SoALayout.DETAILED, None, None, True),
        ("week", SoALayout.DETAILED, True, None, None),
        ("week", SoALayout.DETAILED, None, True, None),
        ("week", SoALayout.DETAILED, None, None, True),
        ("day", SoALayout.DETAILED, True, None, None),
        ("day", SoALayout.DETAILED, None, True, None),
        ("day", SoALayout.PROTOCOL, True, None, True),
        ("week", SoALayout.OPERATIONAL, True, None, None),
        ("week", SoALayout.OPERATIONAL, True, True, None),
        ("week", SoALayout.OPERATIONAL, True, True, True),
    ],
)
def test_flowchart_html_debug(
    soa_test_data: SoATestData,
    api_client: TestClient,
    time_unit: str,
    layout: SoALayout,
    debug_uids: bool,
    debug_coordinates: bool,
    debug_propagation: bool,
):
    """
    Tests HTML SoA with debugging parameters return a valid HTML table
    """

    # Query parameters
    params = {
        "time_unit": time_unit,
        "layout": layout.value,
        "debug_uids": debug_uids,
        "debug_coordinates": debug_coordinates,
        "debug_propagation": debug_propagation,
    }
    params = {k: v for k, v in params.items() if v is not None}

    # Fetch HTML SoA
    response = api_client.get(
        f"/studies/{soa_test_data.study.uid}/flowchart.html", params=params
    )
    assert_response_status_code(response, 200)
    # THEN returns HTML content-type HTTP header
    assert_response_content_type(response, "text/html")

    # THEN body is a parsable HTML document
    # html.parser doesn't supplement the document with additional elements for the sake of well-formedness
    doc = BeautifulSoup(response.text, features="html.parser")
    # THEN document has HTML tag
    assert doc.find("html"), "<HTML> tag not found"

    # THEN document has TABLE
    assert (table := doc.select_one("html > body > table")), "<TABLE> tag not found"

    # THEN the table rows
    assert table.find_all("tr"), "<TABLE> has no <TR> rows"


@pytest.mark.parametrize(
    "path, time_unit",
    [
        ("/studies/{study_uid}/flowchart", "days"),
        ("/studies/{study_uid}/flowchart", "weeks"),
        ("/studies/{study_uid}/flowchart", "foo"),
        ("/studies/{study_uid}/flowchart.html", "bar"),
        ("/studies/{study_uid}/flowchart.html", "weeks"),
        ("/studies/{study_uid}/flowchart.docx", "days"),
        ("/studies/{study_uid}/flowchart.docx", "weekz"),
    ],
)
def test_endpoints_with_invalid_time_unit(
    api_client: TestClient, soa_test_data: SoATestData, path, time_unit: str | None
):
    response = api_client.get(
        path.format_map({"study_uid": soa_test_data.study.uid}),
        params={"time_unit": time_unit},
    )
    assert_response_status_code(response, 422)
    assert_json_response(response)
    payload = response.json()
    assert (detail := payload["detail"][0])
    assert detail["msg"]
    assert detail["loc"] == ["query", "time_unit"]


@pytest.mark.parametrize(
    "path, data_format, column_headers, soa_preferences",
    [
        (
            "/studies/{study_uid}/detailed-soa-exports",
            "text/csv",
            DETAILED_SOA_EXPORT_COLUMN_HEADERS,
            [False, False],
        ),
        (
            "/studies/{study_uid}/detailed-soa-exports",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            DETAILED_SOA_EXPORT_COLUMN_HEADERS,
            [False, True],
        ),
        (
            "/studies/{study_uid}/detailed-soa-exports",
            "text/xml",
            DETAILED_SOA_EXPORT_COLUMN_HEADERS,
            [True, False],
        ),
        (
            "/studies/{study_uid}/detailed-soa-exports",
            "application/json",
            DETAILED_SOA_JSON_EXPORT_COLUMN_HEADERS,
            [True, True],
        ),
        (
            "/studies/{study_uid}/protocol-soa-exports",
            "text/csv",
            PROTOCOL_SOA_EXPORT_COLUMN_HEADERS,
            [True, True],
        ),
        (
            "/studies/{study_uid}/protocol-soa-exports",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            PROTOCOL_SOA_EXPORT_COLUMN_HEADERS,
            [True, False],
        ),
        (
            "/studies/{study_uid}/protocol-soa-exports",
            "text/xml",
            PROTOCOL_SOA_EXPORT_COLUMN_HEADERS,
            [False, True],
        ),
        (
            "/studies/{study_uid}/protocol-soa-exports",
            "application/json",
            # for JSON output, the exported properties are not filtered
            DETAILED_SOA_JSON_EXPORT_COLUMN_HEADERS,
            [False, False],
        ),
        (
            "/studies/{study_uid}/operational-soa-exports",
            "text/csv",
            OPERATIONAL_SOA_EXPORT_COLUMN_HEADERS,
            [],
        ),
        (
            "/studies/{study_uid}/operational-soa-exports",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            OPERATIONAL_SOA_EXPORT_COLUMN_HEADERS_XLSX,
            [],
        ),
        (
            "/studies/{study_uid}/operational-soa-exports",
            "text/xml",
            OPERATIONAL_SOA_EXPORT_COLUMN_HEADERS,
            [],
        ),
        (
            "/studies/{study_uid}/operational-soa-exports",
            "application/json",
            OPERATIONAL_SOA_JSON_EXPORT_COLUMN_HEADERS,
            [],
        ),
    ],
)
def test_soa_exports(
    api_client: TestClient,
    soa_test_data: SoATestData,
    path: str,
    data_format: str,
    column_headers,
    soa_preferences,
):
    """Test the export endpoints return the expected data format"""
    expected_column_headers = column_headers.copy()
    if soa_preferences:
        show_epochs, show_milestones = soa_preferences
        response = api_client.patch(
            f"/studies/{soa_test_data.study.uid}/soa-preferences",
            json={"show_epochs": show_epochs, "show_milestones": show_milestones},
        )
        assert_response_status_code(response, 200)
        res = response.json()
        assert res["show_epochs"] == show_epochs
        assert res["show_milestones"] == show_milestones
        if show_epochs:
            expected_column_headers.append("epoch")
        if show_milestones:
            expected_column_headers.append("milestone")

    response = api_client.get(
        path.format_map({"study_uid": soa_test_data.study.uid}),
        headers={"Accept": data_format},
    )
    assert_response_status_code(response, 200)
    # THEN returns the expected content-type HTTP header
    assert_response_content_type(response, data_format)

    if data_format == "application/json":
        # THEN payload is valid JSON data
        data = response.json()
        assert data, "empty JSON data"
        assert isinstance(data, list), "unexpected JSON data type"
        # THEN JSON list items have expected properties
        expected_column_headers = set(expected_column_headers)
        for i, item in enumerate(data):
            property_keys = set(item.keys())
            assert (
                property_keys == expected_column_headers
            ), f"Item property keys mismatch in record {i}"

    if data_format == "text/csv":
        # THEN payload is valid CSV data
        reader = csv.reader(response.iter_lines(), dialect=csv.excel)
        num_cols = None
        for line in reader:
            assert line
            if num_cols is None:  # first column
                # THEN column headers in first row have expected text and order
                assert line == expected_column_headers, "Column headers mismatch"
                num_cols = len(line)
            else:
                # THEN all rows have the same number of columns
                assert len(line) == num_cols
        # THEN CSV has rows and columns
        assert num_cols, "empty CSV data"

    if data_format == "text/xml":
        # THEN payload is a valid XML document
        assert (tree := lxml.etree.parse(BytesIO(response.content)))
        assert (root := tree.getroot()), "root tag not found"
        # THEN root XML tag is "items"
        assert root.tag == "items", "root tag is not <items>"
        # THEN data contains records
        assert len(root), "root tag <items> has no children"
        # THEN records have expected properties
        expected_column_headers = set(expected_column_headers)
        property_keys = {child.tag for child in root[0]}
        assert property_keys == expected_column_headers, "Item property keys mismatch"

    if (
        data_format
        == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    ):
        # THEN payload is a valid XLSX document
        workbook = openpyxl.load_workbook(BytesIO(response.content))
        # THEN document has a worksheet
        assert len(workbook.worksheets), "XLSX has no worksheets"
        # THEN worksheet has columns
        worksheet = workbook.worksheets[0]
        assert sum(1 for _ in worksheet.columns), "worksheet 0 has no columns"
        # THEN column headers in first row have expected text and order
        column_headers = [column[0].value for column in worksheet.columns]
        assert column_headers == expected_column_headers, "Column headers mismatch"
        # THEN worksheet has more than 1 row
        num_rows = sum(1 for _ in worksheet.rows)
        assert num_rows > 1, f"worksheet 0 has only {num_rows} rows"


def test_get_study_flowchart_versioned(api_client, soa_test_data):
    """Test study flowchart endpoint with versioning for Protocol SoA

    SCENARIO: A study with SoA data gets locked and unlocked, and SoA of the latest draft version updated.
    Check that the protocol SoA of the locked study version is not modified, but the recent SoA of the latest draft
    version differs.
    Because the SoA snapshot is saved when a study version gets locked,
    this also tests that the SoA snapshot is retrievable.
    """

    # GIVEN: a study with protocol SoA
    response = api_client.get(
        f"/studies/{soa_test_data.study.uid}/flowchart",
        params={"layout": SoALayout.PROTOCOL.value},
    )
    assert_response_status_code(response, 200)
    initial_soa = response.json()
    assert len(initial_soa["rows"]) > initial_soa["num_header_rows"]

    # GIVEN: study has a locked version

    response = api_client.patch(
        f"/studies/{soa_test_data.study.uid}",
        json=StudyPatchRequestJsonModel(
            current_metadata=StudyMetadataJsonModel(
                study_description=StudyDescriptionJsonModel(
                    study_title=TestUtils.random_str(prefix="Title ")
                )
            )
        ).model_dump(),
    )
    assert_response_status_code(response, 200)

    response = api_client.post(
        f"/studies/{soa_test_data.study.uid}/locks",
        json=StatusChangeDescription(change_description="Locking good").model_dump(),
    )
    assert_response_status_code(response, 201)

    locked_study_version = response.json()["current_metadata"]["version_metadata"][
        "version_number"
    ]

    response = api_client.delete(f"/studies/{soa_test_data.study.uid}/locks")
    assert_response_status_code(response, 200)

    # GIVEN: SoA has some alterations
    sched = soa_test_data.study_activity_schedules[-1]
    response = api_client.delete(
        f"/studies/{soa_test_data.study.uid}/study-activity-schedules/{sched.study_activity_schedule_uid}"
    )
    assert_response_status_code(response, 204)

    # WHEN: retrieving protocol SoA of draft version
    response = api_client.get(
        f"/studies/{soa_test_data.study.uid}/flowchart",
        params={"layout": SoALayout.PROTOCOL.value},
    )
    assert_response_status_code(response, 200)
    recent_soa = response.json()

    # WHEN: retrieving protocol SoA of the locked version
    response = api_client.get(
        f"/studies/{soa_test_data.study.uid}/flowchart",
        params={
            "layout": SoALayout.PROTOCOL.value,
            "study_value_version": locked_study_version,
        },
    )
    assert_response_status_code(response, 200)
    locked_soa = response.json()

    # THEN: SoA of the locked study version is the same as the initial SoA
    assert locked_soa == initial_soa

    # THEN: SoA of the latest draft study version is different from the SoA of the locked study version
    assert recent_soa != locked_soa
